from docx import Document
from docx.shared import Inches

docx = Document('python-tests.docx')
text_file = open("Output.txt", "w")
s=""
# print (docx)

document = Document()


for p in docx.paragraphs:

    s+=p.text
    s+="\r\n"

text_file.write(s)
text_file.close()


#document.save('output')

# document = docx.read()
# docText = '\n\n'.join([
#     paragraph.text.encode('utf-8') for paragraph in document.paragraphs
# ])
# print (docText)