import lexnlp.extract.en.conditions
import lexnlp.nlp.en.segments.titles
import lexnlp.nlp.en.segments.sections
import lexnlp.nlp.en.segments.paragraphs
import json
import lexnlp.extract.en.citations
import lexnlp.extract.en.definitions
import lexnlp.extract.en.constraints
import lexnlp.extract.en.dates
import lexnlp.extract.en.regulations
import lexnlp.extract.en.urls
from bson import json_util
from flask import Flask, jsonify, request, send_from_directory
from docx import Document
from docx.shared import Inches


app = Flask(__name__, static_url_path='')

@app.route('/')
def index():
   return send_from_directory('static', 'index.html')

#Read Text File
#text = open('testcontract.txt', 'r').read()

@app.route('/uploader', methods = ['GET', 'POST'])
def upload_file():
   if request.method == 'POST':
      f = request.files['file']
      f.save('./uploads/input')
      docx = Document(f)
      text_file = open("uploads/output", "w")
      s=""
      for p in docx.paragraphs:
      	s+=p.text
      	s+="\r\n"
      text_file.write(s)
      text_file.close()
      return 'file uploaded successfully'

@app.route('/contractadvisor/<string:task_id>', methods=['GET'])
def get_task(task_id):
	text = open("uploads/output", 'r').read()

	if len(text)==0:
		return "No File Found"
	
#	if task_id=="title":
#		return getTitle(text)

	if task_id=="suspicious":
		return getSuspicious(text)

	if task_id=="conditions":
		return getConditions(text)

	if task_id=="startdate":
		return getStartDate(text)


# print("Title of the contract")	
# print(list(lexnlp.nlp.en.segments.titles.get_titles(text)))


# #CONDITIONS
# print("Number of conditions")
# x=list(lexnlp.extract.en.conditions.get_conditions(text))
# print(len(x))

# row_json = json.dumps(x)
# print(row_json)

# #CONSTRAINTS
# print("Number of constraints")
# x=list(lexnlp.extract.en.constraints.get_constraints(text))
# print(len(x))

# #DATE
# print(list(lexnlp.extract.en.dates.get_dates(text)))

# #REGULATIONS
# print(list(lexnlp.extract.en.regulations.get_regulations(text)))

# #URLS
# print(list(lexnlp.extract.en.urls.get_urls(text)))


def getTitle(text):
   title=list(lexnlp.nlp.en.segments.titles.get_titles(text))
	raw_json = json.dumps(title)
	return(raw_json)

def getSuspicious(text):
	noOfConditions=len(list(lexnlp.extract.en.conditions.get_conditions(text)))
	if(noOfConditions>25):
		return "Contract looks suspicious as there are around %s conditions" % (noOfConditions)
	else:
		return "Contract is not suspicious as there is only %s" % (noOfConditions)

def getConditions(text):
	conditions=list(lexnlp.extract.en.conditions.get_conditions(text))
	raw_json = json.dumps(conditions)
	return(raw_json)

def getStartDate(text):
	date=list(lexnlp.extract.en.dates.get_dates(text))
	if(len(date)>0):
		return(json.dumps(date[0],default=date_handler))
	else:
		return "No Dates found"

def date_handler(obj):
    return obj.isoformat() if hasattr(obj, 'isoformat') else obj


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0')

